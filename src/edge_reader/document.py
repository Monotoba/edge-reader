from __future__ import annotations

import html
import shutil
import subprocess
import tempfile
from pathlib import Path

from bs4 import BeautifulSoup

from .models import LoadedDocument
from .textseg import normalize_document_text

_TEXT_EXTS = {".txt", ".md", ".markdown", ".rst", ".log", ".csv", ".json", ".xml", ".yaml", ".yml"}
_HTML_EXTS = {".html", ".htm", ".xhtml"}
_CALIBRE_EXTS = {".mobi", ".azw", ".azw3", ".fb2", ".lit", ".pdb"}


class DocumentLoadError(RuntimeError):
    pass


def load_document(path: str | Path) -> LoadedDocument:
    """Load a document into plain text for read-aloud display.

    Supported directly: TXT/Markdown/RST/CSV/JSON/XML/YAML, HTML/XHTML, EPUB, PDF,
    DOCX, RTF. MOBI/AZW/AZW3/FB2/LIT/PDB are supported when Calibre's ebook-convert
    command is installed and discoverable in PATH.
    """
    p = Path(path).expanduser().resolve()
    if not p.exists() or not p.is_file():
        raise DocumentLoadError(f"File does not exist: {p}")

    ext = p.suffix.lower()
    if ext in _TEXT_EXTS:
        text = _read_text_file(p)
    elif ext in _HTML_EXTS:
        text = _html_to_text(_read_text_file(p))
    elif ext == ".epub":
        text = _epub_to_text(p)
    elif ext == ".pdf":
        text = _pdf_to_text(p)
    elif ext == ".docx":
        text = _docx_to_text(p)
    elif ext == ".rtf":
        text = _rtf_to_text(p)
    elif ext in _CALIBRE_EXTS:
        text = _calibre_to_text(p)
    else:
        raise DocumentLoadError(
            f"Unsupported extension {ext!r}. Supported: txt, md, html, epub, pdf, docx, rtf; "
            "mobi/azw3/fb2 require Calibre ebook-convert."
        )

    text = normalize_document_text(text)
    if not text.strip():
        raise DocumentLoadError(f"No readable text found in {p.name}")
    return LoadedDocument(path=p, title=p.stem, text=text)


def _read_text_file(path: Path) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentLoadError(f"Could not decode text file: {path}")


def _html_to_text(markup: str) -> str:
    soup = BeautifulSoup(markup, "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    return html.unescape(soup.get_text("\n"))


def _epub_to_text(path: Path) -> str:
    try:
        from ebooklib import ITEM_DOCUMENT, epub
    except Exception as exc:  # pragma: no cover - dependency/runtime import guard
        raise DocumentLoadError("EPUB support requires ebooklib.") from exc

    try:
        book = epub.read_epub(str(path))
    except Exception as exc:
        raise DocumentLoadError(f"Could not read EPUB: {path.name}") from exc

    parts: list[str] = []
    for item in book.get_items_of_type(ITEM_DOCUMENT):
        content = item.get_content().decode("utf-8", errors="replace")
        text = _html_to_text(content)
        if text.strip():
            parts.append(text)
    return "\n\n".join(parts)


def _pdf_to_text(path: Path) -> str:
    try:
        import fitz  # PyMuPDF
    except Exception as exc:  # pragma: no cover
        raise DocumentLoadError("PDF support requires PyMuPDF.") from exc

    try:
        parts = []
        with fitz.open(path) as doc:
            for page in doc:
                parts.append(page.get_text("text"))
        return "\n\n".join(parts)
    except Exception as exc:
        raise DocumentLoadError(f"Could not extract PDF text: {path.name}") from exc


def _docx_to_text(path: Path) -> str:
    try:
        from docx import Document
    except Exception as exc:  # pragma: no cover
        raise DocumentLoadError("DOCX support requires python-docx.") from exc

    try:
        document = Document(str(path))
        parts = [paragraph.text for paragraph in document.paragraphs]
        # Include readable table cells.
        for table in document.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text for cell in row.cells))
        return "\n".join(parts)
    except Exception as exc:
        raise DocumentLoadError(f"Could not read DOCX: {path.name}") from exc


def _rtf_to_text(path: Path) -> str:
    try:
        from striprtf.striprtf import rtf_to_text
    except Exception as exc:  # pragma: no cover
        raise DocumentLoadError("RTF support requires striprtf.") from exc
    return rtf_to_text(_read_text_file(path))


def _calibre_to_text(path: Path) -> str:
    exe = shutil.which("ebook-convert")
    if not exe:
        raise DocumentLoadError(
            f"{path.suffix.upper()} requires Calibre. Install Calibre and ensure ebook-convert is in PATH."
        )

    with tempfile.TemporaryDirectory(prefix="edge_reader_calibre_") as td:
        out = Path(td) / "converted.txt"
        proc = subprocess.run(
            [exe, str(path), str(out)],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            check=False,
        )
        if proc.returncode != 0 or not out.exists():
            raise DocumentLoadError(
                "Calibre conversion failed.\n"
                f"stdout:\n{proc.stdout[-2000:]}\n\nstderr:\n{proc.stderr[-2000:]}"
            )
        return _read_text_file(out)
